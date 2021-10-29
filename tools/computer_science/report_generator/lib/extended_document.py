from typing import Any
from bs4.element import NavigableString, Tag
import docx
import docx.document
import docx.oxml
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.package import Package
import docx.text
import docx.text.paragraph
import docx.text.run
from docx import Document
from docx.shared import Cm, Emu, Pt
from lxml import etree
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from bs4 import BeautifulSoup
from docx.shared import Pt
import re
import markdown
import docx.api

from lib.utils import dir_name_of_file, real_path_of_file, try_turn_string_to_number, text_to_simple_table
from lib.latex_utils import LatexExtender, latex_to_mathml
from docx.opc.constants import CONTENT_TYPE as CT

from pathlib import Path
import os
import traceback
import json

dir_path = dir_name_of_file(__file__)
MML2OMMLXML_PATH = dir_path+"../res/mml2omml.xsl"

class DocumentConfig:
    font_name:str = 'Times New Roman'
    chapter_title_alignment:str = 'center'
    paragraph_alignment:str = 'justify'
    picture_alignment:str = 'center'
    table_alignment:str = 'center'
    font_size:int = 14
    caption_font_size:int = 14
    caption_line_spacing:float = 1
    table_of_contents_title:str = "Содержание"
    table_of_contents_title_alignment:str = "center"
    line_spacing:float = 1.5
    paragraph_first_line_indent:Cm = Cm(1.25)
    paragraph_left_indent:Cm = Cm(0.0)
    chapter_title_format:str = "{0}  {1} "
    single_formula_alignment:str = "right"

class ExtendedDocument(docx.document.Document):
    def __init__(self, element, part):
        super().__init__(element, part)
        self.xml_namespaces = self.get_xml_namespaces()
        self.config = DocumentConfig()
    
    @staticmethod
    def from_document(doc:docx.document.Document):
        ext_doc = ExtendedDocument(doc.element, doc.part)
        return ext_doc

    def get_xml_namespaces(self):
        return 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '\
            + 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        
    def latex_to_word(self, latex_input, font_size=14, label=""):
        font_size = font_size*2

        mathml = latex_to_mathml(latex_input, label)
        tree = etree.fromstring(mathml)
        xslt = etree.parse(MML2OMMLXML_PATH)
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)
        root = new_dom.getroot()

        prefix_map = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}

        # improving formula with label
        for run in root.findall('.//m:eqArr', prefix_map):
            params = etree.fromstring(f"""
                            <m:eqArrPr {self.xml_namespaces}>
                                <m:maxDist m:val="1"/>
                                <m:ctrlPr>
                                    <w:rPr>
                                        <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>
                                        <w:i/>
                                        <w:szCs w:val="{font_size}"/>
                                    </w:rPr>
                                </m:ctrlPr>
                            </m:eqArrPr>""")
            run.insert(0, params)

        # applying font size
        for run in root.findall('.//m:sSub', prefix_map):
            params = etree.fromstring(f"""
                            <m:sSubPr {self.xml_namespaces}>
                                <m:ctrlPr>
                                    <w:rPr>
                                        <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>
                                        <w:sz w:val="{font_size}"/>
                                        <w:szCs w:val="{font_size}"/>
                                    </w:rPr>
                                </m:ctrlPr>
                            </m:sSubPr>""")
            run.insert(0, params)
        
        for run in root.findall('.//m:r', prefix_map):
            params = etree.fromstring(f"""
                            <w:rPr {self.xml_namespaces}>
								<w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>
								<w:sz w:val="{font_size}"/>
								<w:szCs w:val="{font_size}"/>
							</w:rPr>""")
            run.insert(0, params)
        
        return root

    def make_paragraph(self, document:docx.document.Document, align="left"):
        pg = document.add_paragraph()
        pg.paragraph_format.left_indent = self.config.paragraph_left_indent
        pg.paragraph_format.first_line_indent = self.config.paragraph_first_line_indent
        pg.paragraph_format.line_spacing = self.config.line_spacing
        pg.alignment = self.get_paragraph_alignment(align)
        if align=="center":
            pg.paragraph_format.first_line_indent = pg.paragraph_format.left_indent
            pg.paragraph_format.alignment = self.get_paragraph_alignment("center")
        return pg

    def make_paragraph_text_run(self, paragraph:docx.text.paragraph.Paragraph, text:str, font_size=14):
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = self.config.font_name
        return run

    def get_paragraph_alignment(self, value=None):
        enum = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE}
        if value in enum:
            return enum[value]
        return enum[self.config.paragraph_alignment]

    def delete_paragraph(self, paragraph:docx.text.paragraph.Paragraph):
        pg = paragraph._element
        pg.getparent().remove(pg)
        pg._p = pg._element = None

    def add_hyperlink(self, paragraph, text, url):
        # took from here https://github.com/python-openxml/python-docx/issues/384#issuecomment-294853130
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink

    def add_bookmark(self, paragraph, bookmark_text, bookmark_name, font_size=14) -> docx.text.run.Run:
        run = paragraph.add_run()
        run.font.size = Pt(font_size)
        run.font.name = self.config.font_name
        tag = run._r  # for reference the following also works: tag =  document.element.xpath('//w:r')[-1]
        start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
        start.set(docx.oxml.ns.qn('w:id'), '0')
        start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(start)

        text = docx.oxml.OxmlElement('w:r')
        text.text = bookmark_text
        tag.append(text)

        end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
        end.set(docx.oxml.ns.qn('w:id'), '0')
        end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(end)   
        return run

    def add_bookmark_page_reference(self, paragraph:docx.text.paragraph.Paragraph, bookmark_name:str):
        # run = paragraph.add_run()
        # tag = run._r

        start_r = docx.oxml.OxmlElement('w:r')
        start = docx.oxml.shared.OxmlElement('w:fldChar')
        start.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
        start_r.append(start)
        paragraph._p.append(start_r)

        instr_text_r = docx.oxml.OxmlElement('w:r')
        instr_text = docx.oxml.shared.OxmlElement('w:instrText')
        instr_text.set(docx.oxml.ns.qn('xml:space'), 'preserve')
        instr_text.text = f"PAGEREF {bookmark_name} \\h"
        instr_text_r.append(instr_text)
        paragraph._p.append(instr_text_r)

        separate_r = docx.oxml.OxmlElement('w:r')
        separate = docx.oxml.shared.OxmlElement('w:fldChar')
        separate.set(docx.oxml.ns.qn('w:fldCharType'), 'separate')
        separate_r.append(separate)
        paragraph._p.append(separate_r)

        text_r = docx.oxml.OxmlElement('w:r')
        rpr = docx.oxml.shared.OxmlElement('w:rPr')
        noProof = docx.oxml.shared.OxmlElement('w:noProof')
        rpr.append(noProof)
        text = docx.oxml.shared.OxmlElement('w:t')
        text.text = "1"
        text_r.append(noProof)
        text_r.append(text)
        paragraph._p.append(text_r)

        end_r = docx.oxml.OxmlElement('w:r')
        end = docx.oxml.shared.OxmlElement('w:fldChar')
        end.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
        end_r.append(end)
        paragraph._p.append(end_r)

    def add_ptab(self, paragraph:docx.text.paragraph.Paragraph):
        run = paragraph.add_run()
        ptab_el = etree.fromstring(
            f'<w:ptab {self.xml_namespaces} w:alignment="right" w:leader="dot" w:relativeTo="margin"/>')
        run._element.append(ptab_el)
        return run

    def replace_regex(self, doc, regex, replace):
        for p in doc.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_regex(cell, regex, replace)