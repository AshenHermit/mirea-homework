from typing import Any
from bs4.element import NavigableString, Tag
import docx
import docx.document
import docx.oxml
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.package import Package
import docx.text
import docx.text.paragraph
import docx.text.run
from docx import Document
from docx.shared import Cm, Emu, Pt
from lxml import etree
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from bs4 import BeautifulSoup
from docx.shared import Pt
import re
import markdown
import docx.api
from lib.extended_document import ExtendedDocument
from lib.image_provider import ImageProvider

from lib.utils import dir_name_of_file, real_path_of_file, try_turn_string_to_number, text_to_simple_table
from lib.latex_utils import LatexExtender, latex_to_mathml
from docx.opc.constants import CONTENT_TYPE as CT

from pathlib import Path
import os
import traceback
import json
dir_path = dir_name_of_file(__file__)

class DocxReportRenderer:
    def __init__(self, title_page_filepath="title_page.docx", vars_html_filepath="", image_magick_convert_cmd="convert") -> None:
        self.title_page_filepath = title_page_filepath
        self.html_file_directory = ""
        self.image_provider = ImageProvider("", image_magick_convert_cmd)
        self.soup: BeautifulSoup = None
        self.document: ExtendedDocument = self.create_new_document()

        self.shared_vars = {}
        self.load_vars_from_html(vars_html_filepath)
    
    def load_vars_from_html(self, filepath=""):
        if filepath == "": return
        file = Path(filepath)
        if not file.exists():
            print(f"file {filepath} does not exist")
            return
        html = file.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, 'html.parser')
        self.scan_for_var_tags(soup)
    
    def get_shared_var(self, key:str, default:Any):
        if key in self.shared_vars:
            return self.shared_vars[key]
        else:
            return default
    def get_tag_attr(self, tag:Tag, key:str, default:Any):
        return try_turn_string_to_number(tag.get(key, default))
    def add_shared_var(self, key:str, value:str):
        self.shared_vars[key] = try_turn_string_to_number(value)

    def create_new_document(self) -> ExtendedDocument:
        if self.title_page_filepath == "":
            doc = ExtendedDocument.from_document(docx.Document())
        else:
            doc:docx.document.Document = ExtendedDocument.from_document(docx.Document(self.title_page_filepath))
            doc.add_page_break()
            return doc
        
    def open_html_file(self, filepath):
        file = Path(filepath)
        if file.exists():
            self.load_html(file.read_text(encoding='utf-8'))
            self.html_file_directory = dir_name_of_file(filepath)
            self.image_provider.set_working_directory(self.html_file_directory)
        else:
            self.load_html("")
            raise FileNotFoundError()

    def scan_for_var_tags(self, soup:BeautifulSoup):
        var_tags = soup.find_all("var")
        for var in var_tags:
            var:Tag = var
            if var.has_attr("name"): 
                self.add_shared_var(var.get("name"), var.decode_contents())
    def apply_shared_vars(self, html:str):
        for key in self.shared_vars.keys():
            regex = re.compile(r"(?![a-zA-Z])\$"+key+r"(?![a-zA-Z])")
            value = str(self.shared_vars[key])
            html = regex.sub(value, html)
        return html

    def load_html(self, html=""):
        self.html = html
        self.html = self.html.replace("\n", "")
        self.soup = BeautifulSoup(self.html, 'html.parser')
        # apply variables
        self.scan_for_var_tags(self.soup)
        self.html = self.apply_shared_vars(self.html)
        self.soup = BeautifulSoup(self.html, 'html.parser')
    
    def write_title_page(self):
        title_page = etree.parse(self.title_page_filepath)
        for el in title_page.findall("./"):
            self.document._body._element.append(el)
        self.document.add_page_break()

    def apply_tag_style_to_run(self, run:docx.text.run.Run, tag:Tag):
        if tag.name == "b":
            run.bold = True
        if tag.name == "i":
            run.italic = True

    def add_formula_tag_to_paragraph(self, paragraph:docx.text.paragraph.Paragraph, f_tag:Tag, font_size=14):
        latex_text = LatexExtender(f_tag.get("flags", "")).extend_latex(f_tag.text)
        label = f_tag.get("label", "")
        word_math = self.document.latex_to_word(latex_text, font_size, label)
        paragraph._element.append(word_math)

    def process_paragraph_child(self, paragraph:docx.text.paragraph.Paragraph, tag:Tag, font_size=14, index=1):
        if type(tag) is NavigableString or type(tag) is str:
            # just a text
            if type(tag) is NavigableString: 
                if index==0: text = tag.lstrip()
                else: text = str(tag)
            if type(tag) is str: text = tag
            run = self.document.make_paragraph_text_run(paragraph, text, font_size)
        elif type(tag) is Tag:
            # <f> tags
            if tag.name=="f":
                self.add_formula_tag_to_paragraph(paragraph, tag, font_size)
            # <f> tags
            elif tag.name=="a":
                self.document.add_hyperlink(paragraph, tag.text, tag.get("href", ""))
            # <b>, <i>... tags
            else:
                text = tag.text
                run = self.document.make_paragraph_text_run(paragraph, text, font_size)
                self.apply_tag_style_to_run(run, tag)

    def render_tag_to_paragraph(self, paragraph:docx.text.paragraph.Paragraph, tag:Tag):
        font_size = self.get_tag_attr(tag, "font-size", self.document.config.font_size)

        for i,child in enumerate(tag.children):
            self.process_paragraph_child(paragraph, child, index=i)

    def add_paragraph_tag(self, pg_tag:Tag):
        align = pg_tag.get("align", self.document.config.paragraph_alignment)
        children = list(pg_tag.children)
        if len(children)==1:
            if children[0].name == "f":
                align = self.document.config.single_formula_alignment
        
        pg = self.document.make_paragraph(self.document, align)
        self.render_tag_to_paragraph(pg, pg_tag)
        return pg

    def add_table(self, table_tag:Tag):
        # caption
        self.document.make_paragraph(self.document)
        caption = table_tag.get("caption", "")
        if caption!="":
            pg = self.document.make_paragraph(self.document, "center")
            run = pg.add_run(caption)
            run.font.size = Pt(self.document.config.caption_font_size)

        table_data = text_to_simple_table(table_tag.renderContents().decode("utf-8"))
        cell_width = self.get_tag_attr(table_tag, "cell-width", 10)
        cell_height = self.get_tag_attr(table_tag, "cell-height", 10)

        #TODO: add size checking. actually better to do a class for table
        table = self.document.add_table(len(table_data), len(table_data[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for r,row in enumerate(table.rows):
            row.height = Pt(cell_height)
            for c,cell in enumerate(row.cells):
                cell.width = Pt(cell_width)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self.document.delete_paragraph(cell.paragraphs[0])

                pg = self.document.make_paragraph(cell, "center")
                tag = BeautifulSoup(table_data[r][c], "html.parser")
                children = list(tag.children)
                if len(children)>0:
                    tag = children[0]
                    self.process_paragraph_child(pg, tag, self.document.config.font_size)

    def add_picture(self, img_tag:Tag):
        pg = self.document.make_paragraph(self.document, img_tag.get("align", self.document.config.picture_alignment))

        run = pg.add_run()

        width = self.get_tag_attr(img_tag, "width", None)
        if not width is None: width = Cm(int(width))
        height = self.get_tag_attr(img_tag, "height", None)
        if not height is None: height = Cm(int(height))

        filepath = self.get_tag_attr(img_tag, "src", "")
        if filepath!="":
            filepath = self.image_provider.prepare_image(filepath)
            run.add_picture(filepath, width=width, height=height)
        caption = img_tag.get("caption", "")
        if caption!="":
            run.add_break(WD_BREAK.LINE)
            run.add_text(caption)
            run.font.size = Pt(self.document.config.caption_font_size)
        
        self.document.make_paragraph(self.document)
        return pg

    def get_bookmark_name(self, text:str):
        text = text.lower()
        text = text.replace("\n", " ").replace("\t", "_").replace(" ", "_")
        text = text[:30]
        return text
    def add_chapter(self, ch_el:Tag, index=0):
        # title
        pg = self.document.make_paragraph(self.document, self.document.config.chapter_title_alignment)
        title_text = self.document.config.chapter_title_format.format(index+1, ch_el["title"].upper())
        bookmark_name = self.get_bookmark_name(title_text)
        title = self.document.add_bookmark(pg, title_text, bookmark_name, self.document.config.font_size)
        title.bold = True
        
        paragraphs = ch_el.findChildren()
        for tag in paragraphs:
            # <p> tags
            if tag.name == "p":
                self.add_paragraph_tag(tag)
            # <img> tags
            elif tag.name == "img":
                self.add_picture(tag)
            # <table> tags
            elif tag.name == "table":
                self.add_table(tag)
        self.document.make_paragraph(self.document)

    def add_table_of_contents(self, chapter_names):
        # title
        pg = self.document.make_paragraph(self.document, self.document.config.table_of_contents_title_alignment)
        title = self.document.make_paragraph_text_run(pg, self.document.config.table_of_contents_title.upper())
        title.bold = True

        # list
        pg = self.document.make_paragraph(self.document)
        pg.paragraph_format.first_line_indent = self.document.config.paragraph_left_indent

        for i,name in enumerate(chapter_names):
            title = self.document.config.chapter_title_format.format(i+1, name.upper())
            pg.add_run(title)
            self.document.add_ptab(pg)
            pg.add_run(" ")
            self.document.add_bookmark_page_reference(pg, self.get_bookmark_name(title))
            pg.add_run().add_break(WD_BREAK.LINE)
        
        self.document.add_page_break()

    def replace_patterns_in_document(self, patterns:dict):
        for key in list(patterns.keys()):
            if not type(key) is str: continue
            
            regex = re.compile(r"attr_" + key)
            replace = patterns[key]
            self.document.replace_regex(self.document, regex, replace)

    def create_docx_from_html(self):
        # <report> tag
        report_tag = self.soup.find("report")
        self.replace_patterns_in_document(report_tag.attrs)
        
        chapters:list[Tag] = report_tag.findChildren("chapter")
        self.chapter_names = [ch_el.get("title", "...") for ch_el in chapters]
        self.add_table_of_contents(self.chapter_names)

        children:list[Tag] = report_tag.findChildren()

        for ch_el in children:
            # <chapter> tags
            if ch_el.name == "chapter":
                self.add_chapter(ch_el, chapters.index(ch_el))
            # <pbr> tags
            if ch_el.name == "pbr":
                self.document.add_page_break()
        
        return self.document

    def convert_to_docx(self, output_filepath=""):
        try:
            if output_filepath=="": raise Exception("filepath cant be empty")
            document = self.create_docx_from_html()
            document.save(output_filepath)
            print()
            print("document successfully generated, see file")
            print(real_path_of_file(output_filepath))

        except Exception as e:
            traceback.print_exc()
            print("something went wrong")
        
