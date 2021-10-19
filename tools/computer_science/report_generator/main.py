import docx
import docx.document
import docx.text
import docx.text.paragraph
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from lxml import etree

from bs4 import BeautifulSoup
import latex2mathml.converter
from docx.shared import Pt

import markdown

from pathlib import Path

class HtmlReportToDocx:
    def __init__(self) -> None:
        self.soup: BeautifulSoup = None
        
    def open_file(self, filepath):
        file = Path(filepath)
        if file.exists():
            self.load_html(file.read_text())
        else:
            self.load_html("")
            raise FileNotFoundError()
            
    def load_html(self, html=""):
        self.html = html
        self.html = self.html.replace("\n", "")
        self.soup = BeautifulSoup(self.html, 'html.parser')

    def latex_to_word(self, latex_input):
        mathml = latex2mathml.converter.convert(latex_input)
        tree = etree.fromstring(mathml)
        xslt = etree.parse(
            'MML2OMML.XSL'
            )
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)
        return new_dom.getroot()

    def create_docx_from_html(self):
        root_el = self.soup.find("report")

        paragraphs:list[docx.text.paragraph.Paragraph] = self.soup.find("report").findChildren("p")
        for p in paragraphs:
            p["title"]
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = Cm(1.2)
            p.add_run("awdawd awd aw wdagegaeg").font.size = Cm(1.2)

        document = docx.Document()

    def convert_to_docx(self, output_filepath=None):
        pass
        

        
        
        

document:docx.document.Document = docx.Document()
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')